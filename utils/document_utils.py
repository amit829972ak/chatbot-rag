import os
import io
import pandas as pd
from typing import Tuple, Optional
from PIL import Image

def process_document(file, file_type: str) -> Tuple[str, Optional[pd.DataFrame]]:
    """
    Process various document types and extract their content.
    
    Args:
        file: The uploaded file object
        file_type (str): The file extension/type
        
    Returns:
        Tuple[str, Optional[pd.DataFrame]]: Tuple containing extracted text content and dataframe (if applicable)
    """
    file_type = file_type.lower()
    df = None
    content = ""
    
    try:
        # CSV files
        if file_type == '.csv':
            df = pd.read_csv(file)
            content = convert_df_to_csv_string(df)
        
        # TSV files
        elif file_type == '.tsv':
            df = pd.read_csv(file, sep='\t')
            content = convert_df_to_csv_string(df)
        
        # Excel files
        elif file_type == '.xlsx' or file_type == '.xls':
            df = pd.read_excel(file)
            content = convert_df_to_csv_string(df)
        
        # PDF files
        elif file_type == '.pdf':
            # Use PyPDF2 to extract text from PDFs
            try:
                import PyPDF2
                from PyPDF2 import PdfReader
                
                reader = PdfReader(file)
                content = ""
                
                # Extract text from each page
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        content += f"--- Page {i+1} ---\n{page_text}\n\n"
                
                if not content.strip():
                    content = "[PDF contains no extractable text content. It may be scanned or contain only images.]"
            except Exception as e:
                content = f"[Error extracting PDF content: {str(e)}]"
        
        # Text files
        elif file_type == '.txt':
            content = file.getvalue().decode('utf-8')
        
        # Word documents
        elif file_type == '.docx' or file_type == '.doc':
            try:
                from docx import Document
                document = Document(file)
                
                # Extract text from paragraphs
                paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
                content = "\n\n".join(paragraphs)
                
                # Extract text from tables
                if document.tables:
                    content += "\n\n--- Tables ---\n"
                    for i, table in enumerate(document.tables):
                        content += f"\nTable {i+1}:\n"
                        for row in table.rows:
                            row_text = [cell.text for cell in row.cells]
                            content += " | ".join(row_text) + "\n"
            except Exception as e:
                content = f"[Error extracting DOCX content: {str(e)}]"
        
        else:
            content = f"[Unsupported file type: {file_type}]"
        
        return content, df
        
    except Exception as e:
        error_msg = f"Error processing document: {str(e)}"
        return error_msg, None

def get_file_extension(filename: str) -> str:
    """
    Get the lowercase file extension including the dot.
    
    Args:
        filename (str): The filename
        
    Returns:
        str: Lowercase file extension with dot
    """
    _, ext = os.path.splitext(filename)
    return ext.lower()

def convert_df_to_csv_string(df: pd.DataFrame) -> str:
    """
    Convert a DataFrame to a CSV string.
    
    Args:
        df (pd.DataFrame): DataFrame to convert
        
    Returns:
        str: CSV string representation
    """
    try:
        # Get basic DataFrame info
        rows, cols = df.shape
        info = f"DataFrame: {rows} rows × {cols} columns\n\n"
        
        # Add column types
        info += "Column Types:\n"
        for col, dtype in df.dtypes.items():
            info += f"- {col}: {dtype}\n"
        info += "\n"
        
        # Get DataFrame summary
        info += "Summary Statistics:\n"
        try:
            summary = df.describe(include='all').to_string()
            info += summary + "\n\n"
        except:
            info += "[Unable to generate summary statistics]\n\n"
        
        # First 100 rows as CSV (with limit for very large DataFrames)
        limit = 100 if len(df) > 100 else len(df)
        info += f"First {limit} rows:\n"
        csv_text = df.head(limit).to_csv(index=False)
        
        # Combine everything
        final_content = info + csv_text
        
        return final_content
    except Exception as e:
        return f"Error converting DataFrame to string: {str(e)}"

def convert_df_to_json_string(df: pd.DataFrame) -> str:
    """
    Convert a DataFrame to a JSON string.
    
    Args:
        df (pd.DataFrame): DataFrame to convert
        
    Returns:
        str: JSON string representation
    """
    try:
        # Limit to first 100 rows for very large DataFrames
        limit = 100 if len(df) > 100 else len(df)
        sample_df = df.head(limit)
        
        # Convert to JSON
        json_str = sample_df.to_json(orient='records', indent=2)
        
        # Add metadata about truncation
        if len(df) > limit:
            return f"DataFrame: {len(df)} rows × {len(df.columns)} columns (showing first {limit} rows)\n\n{json_str}"
        else:
            return f"DataFrame: {len(df)} rows × {len(df.columns)} columns\n\n{json_str}"
    except Exception as e:
        return f"Error converting DataFrame to JSON: {str(e)}"

def get_document_summary(document_content: str) -> str:
    """
    Create a brief summary of the document content.
    
    Args:
        document_content (str): The document content
        
    Returns:
        str: A brief summary
    """
    # Simple summary based on content length
    words = document_content.split()
    word_count = len(words)
    
    if word_count == 0:
        return "Empty document"
    
    char_count = len(document_content)
    first_line = document_content.split('\n')[0][:100] + "..."
    
    summary = f"Document ({word_count} words, {char_count} characters) beginning with: {first_line}"
    
    return summary
